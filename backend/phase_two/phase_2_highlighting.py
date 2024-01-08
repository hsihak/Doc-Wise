# Import necessary libraries for semantic text similarity, document processing, and utility functions
from semantic_text_similarity.models import WebBertSimilarity, ClinicalBertSimilarity
from docx import Document
from sentence_transformers import SentenceTransformer, util
from sklearn.feature_extraction.text import CountVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import torch
import os
import PyPDF2
from pdf2docx import Converter

class PhaseTwoHighlightDifferences: 
    def __init__(self, model_name='distilbert-base-nli-stsb-mean-tokens'):
        # Load the Sentence-BERT model for semantic similarity analysis of sentences
        self.model = SentenceTransformer(model_name)

    # Function to convert DOC files to DOCX format, as DOCX is easier to process
    def convert_to_docx(doc_file):
        converted_file = doc_file.replace(".doc", ".docx")
        docx_file = Document()
        docx_file.add_paragraph(doc_file)
        docx_file.save(converted_file)
        return converted_file
    
    # Function to identify and highlight differences between two documents.
    # It reads text from each document, computes semantic sentence embeddings using Sentence-BERT,
    # and then identifies sentences with low similarity across the documents based on a specified threshold.
    # The function returns two lists: one with highlighted sentences from Document A, and the other from Document B,
    # allowing for a visual comparison of dissimilar content between the two documents.
    def get_highlighted_sentences(self, file_path_A, file_path_B, threshold):
        file_type_A = file_path_A.split('.')[-1].lower()
        file_type_B = file_path_B.split('.')[-1].lower()

        if file_type_A == 'docx':
            # Load the document from docx file
            doc_A = Document(file_path_A)
            text_A = ' '.join([paragraph.text for paragraph in doc_A.paragraphs])
        elif file_type_A == 'doc':
            # Convert DOC file to DOCX format
            converted_file_A = self.convert_to_docx(file_path_A)
            # Load the text from converted DOCX file using python-docx library
            doc_A = Document(converted_file_A)
            text_A = ' '.join([paragraph.text for paragraph in doc_A.paragraphs])
            # Remove the temporary converted file
            os.remove(converted_file_A)
        elif file_type_A == 'pdf':
            # Load the text from pdf file using a different library
            pdf_reader = PyPDF2.PdfReader(file_path_A)
            text_A = ' '.join([pdf_reader.pages[page_num].extract_text() for page_num in range(len(pdf_reader.pages))])
        else:
            raise ValueError("Unsupported file type: {}".format(file_type_A))

        if file_type_B == 'docx':
            # Load the document from docx file
            doc_B = Document(file_path_B)
            text_B = ' '.join([paragraph.text for paragraph in doc_B.paragraphs])
        elif file_type_B == 'doc':
            # Convert DOC file to DOCX format
            converted_file_B = self.convert_to_docx(file_path_B)
            # Load the text from converted DOCX file using python-docx library
            doc_B = Document(converted_file_B)
            text_B = ' '.join([paragraph.text for paragraph in doc_B.paragraphs])
            # Remove the temporary converted file
            os.remove(converted_file_B)
        elif file_type_B == 'pdf':
            # Load the text from pdf file using a different library
            pdf_reader = PyPDF2.PdfReader(file_path_B)
            text_B = ' '.join([pdf_reader.pages[page_num].extract_text() for page_num in range(len(pdf_reader.pages))])
        else:
            raise ValueError("Unsupported file type: {}".format(file_type_B))

        # Split the text into sentences
        sentences_A = text_A.split('. ')
        sentences_B = text_B.split('. ')

        # Remove empty sentences
        sentences_A = [sentence.strip() for sentence in sentences_A if sentence.strip()]
        sentences_B = [sentence.strip() for sentence in sentences_B if sentence.strip()]

        # Compute sentence embeddings
        embeddings_A = self.model.encode(sentences_A)
        embeddings_B = self.model.encode(sentences_B)

        # list of highlighted sentences in A
        highlighted_sentences_A = []
        all_sentences_A = []

        # Iterate over sentences in document A
        for i, sentence_A in enumerate(sentences_A):
            # Compute cosine similarity between the sentence in A and all sentences in B
            similarity_scores = util.pytorch_cos_sim(torch.tensor(embeddings_A[i]).unsqueeze(0), torch.tensor(embeddings_B))

        # Check if the maximum similarity score is below the threshold
            all_sentences_A.append(sentence_A + ". ")
            if torch.max(similarity_scores) < threshold:
                # Highlight the differing sentence in document A by setting its font color to red
                # highlighted_sentences_A.append({'sentence_text': sentence_A + ". ", 'highlight': 1})
                highlighted_sentences_A.append(sentence_A + ". ")
            # else:
            #     # Add the sentence to the highlighted document without highlighting
            #     highlighted_sentences_A.append({'sentence_text': sentence_A + ". ", 'highlight': 0})



        # list of highlighted sentences in B
        highlighted_sentences_B = []
        all_sentences_B = []
        # Iterate over sentences in document B
        for j, sentence_B in enumerate(sentences_B):
            # Compute cosine similarity between the sentence in B and all sentences in A
            similarity_scores = util.pytorch_cos_sim(torch.tensor(embeddings_B[j]).unsqueeze(0), torch.tensor(embeddings_A))

            # Check if the maximum similarity score is below the threshold
            all_sentences_B.append(sentence_B + ". ")
            if torch.max(similarity_scores) < threshold:
                # Highlight the differing sentence in document B by setting its font color to red
                highlighted_sentences_B.append(sentence_B + ". ")
                # highlighted_sentences_B.append({'sentence_text': sentence_B + ". ", 'highlight': 1})
            # else:
            #     # Add the sentence to the highlighted document without highlighting
            #     highlighted_sentences_B.append({'sentence_text': sentence_B + ". ", 'highlight': 0})
        # print(all_sentences_A)
        return all_sentences_A, all_sentences_B, highlighted_sentences_A, highlighted_sentences_B

    # Function to highlight differences between two documents in DOCX, DOC, or PDF format.
    # It extracts text from each document, compares sentences using Sentence-BERT embeddings,
    # and highlights sentences with low similarity. The function creates two new documents
    # with these highlights and saves them to the specified output paths.

    def highlight_differences(self, file_path_A, file_path_B, output_file_path_A, output_file_path_B, threshold):
        file_type_A = file_path_A.split('.')[-1].lower()
        file_type_B = file_path_B.split('.')[-1].lower()

        if file_type_A == 'docx':
            # Load the document from docx file
            doc_A = Document(file_path_A)
            text_A = ' '.join([paragraph.text for paragraph in doc_A.paragraphs])
        elif file_type_A == 'doc':
            # Convert DOC file to DOCX format
            converted_file_A = self.convert_to_docx(file_path_A)
            # Load the text from converted DOCX file using python-docx library
            doc_A = Document(converted_file_A)
            text_A = ' '.join([paragraph.text for paragraph in doc_A.paragraphs])
            # Remove the temporary converted file
            os.remove(converted_file_A)
        elif file_type_A == 'pdf':
            # Load the text from pdf file using a different library
            pdf_reader = PyPDF2.PdfFileReader(file_path_A)
            text_A = ' '.join([pdf_reader.getPage(page_num).extract_text() for page_num in range(pdf_reader.numPages)])
        else:
            raise ValueError("Unsupported file type: {}".format(file_type_A))

        if file_type_B == 'docx':
            # Load the document from docx file
            doc_B = Document(file_path_B)
            text_B = ' '.join([paragraph.text for paragraph in doc_B.paragraphs])
        elif file_type_B == 'doc':
            # Convert DOC file to DOCX format
            converted_file_B = self.convert_to_docx(file_path_B)
            # Load the text from converted DOCX file using python-docx library
            doc_B = Document(converted_file_B)
            text_B = ' '.join([paragraph.text for paragraph in doc_B.paragraphs])
            # Remove the temporary converted file
            os.remove(converted_file_B)
        elif file_type_B == 'pdf':
            # Load the text from pdf file using a different library
            pdf_reader = PyPDF2.PdfFileReader(file_path_B)
            text_B = ' '.join([pdf_reader.getPage(page_num).extract_text() for page_num in range(pdf_reader.numPages)])
        else:
            raise ValueError("Unsupported file type: {}".format(file_type_B))

        # Split the text into sentences
        sentences_A = text_A.split('. ')
        sentences_B = text_B.split('. ')

        # Remove empty sentences
        sentences_A = [sentence.strip() for sentence in sentences_A if sentence.strip()]
        sentences_B = [sentence.strip() for sentence in sentences_B if sentence.strip()]

        # Compute sentence embeddings
        embeddings_A = self.model.encode(sentences_A)
        embeddings_B = self.model.encode(sentences_B)

        # Create a new document for the highlighted version of document A
        highlighted_doc_A = Document()

        # Iterate over sentences in document A
        for i, sentence_A in enumerate(sentences_A):
            # Compute cosine similarity between the sentence in A and all sentences in B
            similarity_scores = util.pytorch_cos_sim(torch.tensor(embeddings_A[i]).unsqueeze(0), torch.tensor(embeddings_B))

        # Check if the maximum similarity score is below the threshold
            if torch.max(similarity_scores) < threshold:
                # Highlight the differing sentence in document A by setting its font color to red
                run_A = highlighted_doc_A.add_paragraph().add_run(sentence_A)
                font_A = run_A.font
                font_A.color.rgb = RGBColor(255, 0, 0)  # Red color
            else:
                # Add the sentence to the highlighted document without highlighting
                highlighted_doc_A.add_paragraph(sentence_A)


        # Save the highlighted document for document A
        highlighted_doc_A.save(output_file_path_A)


        # Create a new document for the highlighted version of document B
        highlighted_doc_B = Document()

        # Iterate over sentences in document B
        for j, sentence_B in enumerate(sentences_B):
            # Compute cosine similarity between the sentence in B and all sentences in A
            similarity_scores = util.pytorch_cos_sim(torch.tensor(embeddings_B[j]).unsqueeze(0), torch.tensor(embeddings_A))

            # Check if the maximum similarity score is below the threshold
            if torch.max(similarity_scores) < threshold:
                # Highlight the differing sentence in document B by setting its font color to red
                run_B = highlighted_doc_B.add_paragraph().add_run(sentence_B)
                font_B = run_B.font
                font_B.color.rgb = RGBColor(255, 0, 0)  # Red color
            else:
                # Add the sentence to the highlighted document without highlighting
                highlighted_doc_B.add_paragraph(sentence_B)

        # Save the highlighted document for document B
        highlighted_doc_B.save(output_file_path_B)

    # this function converts a pdf to a docx file
    def pdf_to_docx(input_pdf_path, output_docx_path_name):
        pdf_file = input_pdf_path
        docx_file = output_docx_path_name
        cv = Converter(pdf_file)
        cv.convert(docx_file, start=0, end=None)
        cv.close()
        return output_docx_path_name
    

    # Function to retrieve font styling information from a 'run' in a DOCX document.
    # Returns a dictionary with text content, font name, size, and style attributes (bold, italic, underline).
    # Essential for preserving text styling when processing document content.
    def read_font_info(run):
        return {
            'text': run.text,  # Text content of the run
            'font_name': run.font.name,  # Font name
            'font_size': run.font.size or 12,  # Font size with a default of 12
            'bold': run.bold,  # Bold styling
            'italic': run.italic,  # Italic styling
            'underline': run.underline,  # Underline styling
    }

    # Simplifies handling of font attributes by treating only explicit 'True' as True, else False.
    # Useful in determining precise font styling, particularly when some attributes might be unset (None).
    def true_false(input):
        return input is True
    

    # Function to determine if a specific sentence is highlighted.
    # It compares the sentence against a list of highlighted sentences (los),
    # checking for direct matches or if the sentence is a substring of any highlighted sentence.
    def sent_highlighted(sentence, los):
        remove = string.punctuation + string.whitespace
        mapping = {ord(c): None for c in remove}
        # check if it's a substring in the metadata
        if sentence.translate(mapping) in los[0].translate(mapping):
            return True
        # check if it matches any other entry in the highlighted list
        for sent in los:
            # equals or is a substring in the highlighted text
            if sent.translate(mapping) == sentence.translate(mapping) or sentence.translate(mapping) in sent.translate(mapping):
                return True
        return False
    
    # Function to split a block of text into individual sentences.
    # This is particularly useful for processing each sentence separately,
    # like when identifying which sentences to highlight.
    def split_text_into_sentence_list(text):
        sentences = text.split('. ') # need space for cases like "e.g. "
        # Remove empty sentences
        sentences = [sentence.strip() for sentence in sentences if sentence.strip()]
        # should give a list of sentences with the period removed
        return sentences
    
    # Function to output formatted paragraphs into a DOCX file.
    # It applies specific font styles and highlights to the sentences based on the comparison results.
    def output_paragraphs(self, output_file_path, docx, highlighted_sentences):
        output = Document()
        in_metadata = True
        for par in docx.paragraphs:
            if not par.text.strip():
            # if its an empty text (e.g. newline), then skip it
                continue
            # get the font information
            runs = par.runs
            if runs: # check that the runs are not empty
            # non empty, then get the font info of the first run and assume that it carries over through the whole paragraph -- !this assumption may result in unexpected fonts
                font_info = self.read_font_info(runs[0])

            los = self.split_text_into_sentence_list(par.text) #los stands for list_of_sentences
            # if the text ends with a period ("."), then we must be out of the metadata
            if par.text.strip() and in_metadata and par.text.strip()[-1] == ".":
            # then we have our first paragraph, and we are no longer in the metadata
                in_metadata = False

            para = output.add_paragraph()
            for sent in los:
            # get font inf
                out_run = para.add_run(sent)
                out_run.font.name = font_info['font_name']
                out_run.font.size = Pt(font_info['font_size'])
                out_run.bold = self.true_false(font_info['bold'])
                out_run.italic = self.true_false(font_info['italic'])
                out_run.underline = self.true_false(font_info['underline'])
                # make sure to account for all forms of sentences (just in case)
                highlight = False
                # account for both header format and sentence format
                if self.sent_highlighted(sent, highlighted_sentences):
                    # then the run should highlighted
                    out_run.font.color.rgb = RGBColor(255, 0, 0)  # Red color
                    highlight = True
                else:
                    out_run.font.color.rgb = RGBColor(0, 0, 0) # black colour
                    highlight = False

                if len(los) == 1 and sent.strip().endswith(":"):
                    # then it is a title, we need to bold this
                    out_run.bold = True
                    # we alo need to combine the next sentence with the next sentence to properly see if it matches with the highlighted sentence
                elif not in_metadata and not sent.strip().endswith("."):
                    # it is not a title, make sure to end the sentences with a ". "
                    # need to make sure it doesn't add periods to metadata
                    run = para.add_run(". ")
                    if highlight:
                        run.font.color.rgb = RGBColor(255, 0, 0)  # Red color
                    else:
                        run.font.color.rgb = RGBColor(0, 0, 0) # black colour
        output.save(output_file_path)


    # Function to determine the final file paths for processing documents A and B.
    # This function checks the file type of each input document (PDF or DOCX) and converts PDFs to DOCX format if necessary.
    # The conversion is essential because the subsequent text processing and highlighting functionalities are optimized for DOCX format.
    def get_final_paths(self, file_path_A, file_path_B):
        file_type_A = file_path_A.split('.')[-1].lower()
        file_type_B = file_path_B.split('.')[-1].lower()

        # Convert document A to DOCX format if it is a PDF, otherwise, use the original path
        if file_type_A == "pdf":
            final_path_A = self.pdf_to_docx(file_path_A, "A.docx")
        else:
            final_path_A = file_path_A

        # Repeat the same conversion process for document B
        if file_type_B == "pdf":
            final_path_B = self.pdf_to_docx(file_path_B, "B.docx")
        else:
            final_path_B = file_path_B

        return final_path_A, final_path_B
    
    # Main function to orchestrate the highlighting and formatting of documents based on semantic similarity.
    def highlight_and_format(self, file_path_A, file_path_B, output_file_path_A, output_file_path_B, threshold):
        # get the list of sentences that are highlighted
        final_path_A, final_path_B = self.get_final_paths(file_path_A, file_path_B)
        all_sentences_A, all_sentences_B, highlighted_sentences_A, highlighted_sentences_B = get_highlighted_sentences(final_path_A, final_path_B, threshold)
        docx_a = Document(final_path_A)
        docx_b = Document(final_path_B)
        self.output_paragraphs(output_file_path_A, docx_a, highlighted_sentences_A)
        self.output_paragraphs(output_file_path_B, docx_b, highlighted_sentences_B)


def main():
    # Initialize PhaseTwoHighlightDifferences Class
    phasetwo_highlightdifferences = PhaseTwoHighlightDifferences()

    # Define file paths and thresholds 
