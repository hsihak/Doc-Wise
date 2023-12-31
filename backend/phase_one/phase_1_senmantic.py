import os
import textract
import PyPDF2
import subprocess
from docx import Document
from sklearn.feature_extraction.text import TfidfVectorizer
from nltk.corpus import stopwords
from nltk.stem import WordNetLemmatizer
import numpy as np
from sklearn.metrics.pairwise import cosine_similarity
import pandas as pd
import csv
from transformers import BertModel, BertTokenizer
import torch
import nltk

class PhaseOneSemanticSimilarity:
    def __init__(self):
        nltk.download('stopwords')
        nltk.download('wordnet')
        self.model = BertModel.from_pretrained('bert-base-uncased', output_hidden_states=True)
        self.tokenizer = BertTokenizer.from_pretrained('bert-base-uncased')
        self.stop_words = set(stopwords.words('english'))
        self.lemmatizer = WordNetLemmatizer()
    
    def encode(self, text):
        input_ids = torch.tensor(self.tokenizer.encode(text, truncation=True, max_length=512)).unsqueeze(0)  # Batch size 1
        outputs = self.model(input_ids)
        # Use the mean of the last layer's hidden states as the sentence representation.
        sentence_embedding = torch.mean(outputs.last_hidden_state, dim=1).detach().numpy()
        return sentence_embedding
    
    def read_preprocess_files(self, temp_file_paths):
        preprocessed_texts = []
        # Processing each file based on its format (DOCX, DOC, PDF, etc.) and extracting text from it.
        for file_path in temp_file_paths:
            if file_path.endswith(".docx"):
                # Reading DOCX file.
                doc = Document(file_path)
                text = [paragraph.text for paragraph in doc.paragraphs]
                preprocessed_texts.append(" ".join(text))

            elif file_path.endswith(".doc"):
                # Converting DOC file to DOCX format.
                docx_file_path = file_path + "x"
                subprocess.call(['soffice', '--headless', '--convert-to', 'docx', file_path])

                # Reading the converted DOCX file.
                doc = Document(docx_file_path)
                text = [paragraph.text for paragraph in doc.paragraphs]
                preprocessed_texts.append(" ".join(text))

                # Removing temporary .docx file.
                os.remove(docx_file_path)

            elif file_path.endswith(".pdf"):
                # Reading PDF file.
                pdf_reader = PyPDF2.PdfReader(file_path)
                text = ""
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text()
                preprocessed_texts.append(text)

            else:
                # Reading other file types using textract.
                text = textract.process(file_path, method='text').decode('utf-8')
                preprocessed_texts.append(text)
            # Applying preprocessing: removing stopwords and lemmatizing the text.
            # This is done to reduce the text to its most informative components.
            preprocessed_texts = [" ".join([self.lemmatizer.lemmatize(word.lower()) for word in text.split() if word.lower() not in self.stop_words]) for text in preprocessed_texts]

        return preprocessed_texts
    

    def output_to_csv(self, similarity_scores, uploaded_files, output_csv_path):
        # Prepare the data with different similarity thresholds.
        data = [["Threshold", "File Pair's Names"], ["Similarity between files > 90%"], ["Similarity between files > 80%"], ["Similarity between files > 70%"], ["Similarity between files > 60%"], ["Similarity between files > 50%"], ["Similarity between files <= 50%"]]

        for i in range(len(uploaded_files)):
            for j in range(i+1, len(uploaded_files)):
                file1 = uploaded_files[i]
                file2 = uploaded_files[j]
                similarity_score = similarity_scores[i][j]
                similarity_score = round(similarity_score * 100, 2)

                # Categorizing file pairs based on their similarity percentage.
                if similarity_score > 90:
                    data[1].append(f"{file1} and {file2}")
                elif similarity_score > 80:
                    data[2].append(f"{file1} and {file2}")
                elif similarity_score > 70:
                    data[3].append(f"{file1} and {file2}")
                elif similarity_score > 60:
                    data[4].append(f"{file1} and {file2}")
                elif similarity_score > 50:
                    data[5].append(f"{file1} and {file2}")
                else:
                    data[6].append(f"{file1} and {file2}")

        # Writing the similarity data to a CSV file.
        # The CSV file will be useful for front-end applications to display or use the similarity data.
        with open(output_csv_path, 'w', newline='') as csvfile:
            writer = csv.writer(csvfile)
            writer.writerows(data)

    def compute_similarity(self, preprocessed_texts):
        # TF-IDF (Term Frequency-Inverse Document Frequency) is a statistical measure used to evaluate the importance of a word
        # to a document in a collection or corpus. Here, it's used to convert the text to a numerical form.
        vectorizer = TfidfVectorizer()
        tfidf_matrix = vectorizer.fit_transform(preprocessed_texts).toarray()

        # Converting the preprocessed texts into BERT embeddings for semantic analysis.
        bert_embeddings = np.array([self.encode(text)[0] for text in preprocessed_texts])

        # Computing cosine similarity separately for TF-IDF and BERT vectors.
        # Cosine similarity measures the cosine of the angle between two non-zero vectors of an inner product space,
        # which is used as a measure of similarity between two vectors.
        tfidf_similarity_scores = cosine_similarity(tfidf_matrix)
        bert_similarity_scores = cosine_similarity(bert_embeddings)

        # Averaging the similarity scores from both TF-IDF and BERT to get a more comprehensive similarity measure.
        average_similarity_scores = np.mean([tfidf_similarity_scores, bert_similarity_scores], axis=0)

        return average_similarity_scores
    

def main():
    phaseone_semanticsimilarity = PhaseOneSemanticSimilarity()

    temp_dir = "C:/Users/Hangsihak Sin/OneDrive/Documents/School/Doc-Wise/backend/phase_one/static/content/temp_files"
    os.makedirs(temp_dir, exist_ok=True)
    temp_file_paths = [os.path.join(temp_dir, file) for file in os.listdir(temp_dir)
                       if file.endswith(".pdf") or file.endswith(".docx")]
    
    print(temp_file_paths)

    print("Current Working Directory:", os.getcwd())
    
    preprocessed_texts = phaseone_semanticsimilarity.read_preprocess_files(temp_file_paths)
    average_similarity_scores = phaseone_semanticsimilarity.compute_similarity(preprocessed_texts)
    # Displaying the average similarity scores for visualization.
    print("\nAverage Similarity Scores:")
    print(average_similarity_scores)


    # Create a DataFrame for better visualization
    average_similarity_df = pd.DataFrame(average_similarity_scores, index=temp_file_paths, columns=temp_file_paths)

    # Save the dataframe to an Excel file
    average_similarity_df.to_excel('C:/Users/Hangsihak Sin/OneDrive/Documents/School/Doc-Wise/client/src/assets/phase-one/static/similarity_scores.xlsx')

    output_csv_path = "C:/Users/Hangsihak Sin/OneDrive/Documents/School/Doc-Wise/client/src/assets/phase-one/static/threshold_similarity_scores.csv"
    phaseone_semanticsimilarity.output_to_csv(average_similarity_scores, temp_file_paths, output_csv_path)

if __name__ == "__main__":
    main()